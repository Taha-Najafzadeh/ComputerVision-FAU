from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt


BASE_DIR = Path(__file__).resolve().parent
DOCX_PATH = BASE_DIR / "exercise_1_2_discussion.docx"


def set_normal_style(document):
    style = document.styles["Normal"]
    style.font.name = "Arial"
    style.font.size = Pt(11)
    paragraph_format = style.paragraph_format
    paragraph_format.space_after = Pt(6)
    paragraph_format.line_spacing = 1.08


def add_title(document, text):
    paragraph = document.add_paragraph()
    paragraph.style = document.styles["Title"]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = paragraph.add_run(text)
    run.bold = True
    run.font.name = "Arial"
    run.font.size = Pt(22)


def add_heading(document, text):
    paragraph = document.add_heading(text, level=1)
    paragraph.runs[0].font.name = "Arial"
    paragraph.runs[0].font.size = Pt(16)


def add_body(document, text):
    paragraph = document.add_paragraph(text)
    paragraph.paragraph_format.space_after = Pt(6)


def add_bullet(document, text):
    paragraph = document.add_paragraph(text, style="List Bullet")
    paragraph.paragraph_format.space_after = Pt(4)


def add_dimensions_table(document):
    add_body(document, "For the first example, the estimated dimensions are:")

    headers = ["Example", "Height", "Length", "Width"]
    rows = [
        ["example1kinect.mat", "0.1961 m", "0.6331 m", "0.4569 m"],
        ["example2kinect.mat", "0.1947 m", "0.4525 m", "0.3633 m"],
        ["example3kinect.mat", "0.1925 m", "0.4555 m", "0.3615 m"],
        ["example4kinect.mat", "0.1877 m", "0.4448 m", "0.3401 m"],
    ]

    table = document.add_table(rows=1, cols=4)
    table.style = "Table Grid"
    table.autofit = True

    header_cells = table.rows[0].cells
    for index, header in enumerate(headers):
        header_cells[index].text = header
        for paragraph in header_cells[index].paragraphs:
            for run in paragraph.runs:
                run.bold = True

    for row in rows:
        cells = table.add_row().cells
        for index, value in enumerate(row):
            cells[index].text = value

    document.add_paragraph()


document = Document()
set_normal_style(document)

section = document.sections[0]
section.top_margin = Inches(1)
section.bottom_margin = Inches(1)
section.left_margin = Inches(1)
section.right_margin = Inches(1)

add_title(document, "Exercise 1.2: Box Detection Discussion")

add_heading(document, "Results")
add_body(
    document,
    "For the result I first detected the floor plane and then used the remaining points to find the "
    "top of the box. The output figures show the main intermediate results:",
)
add_bullet(document, "amplitude image,")
add_bullet(document, "distance image,")
add_bullet(document, "floor mask,")
add_bullet(document, "box top mask,")
add_bullet(document, "overlay with floor, box top and corners,")
add_bullet(document, "a small 3D point cloud plot.")

add_body(document, "The figures I generated are:")
add_bullet(document, "example1kinect_exercise_1_1.png")
add_bullet(document, "example2kinect_exercise_1_1.png")
add_bullet(document, "example3kinect_exercise_1_1.png")
add_bullet(document, "example4kinect_exercise_1_1.png")

add_dimensions_table(document)
add_body(
    document,
    "The height is calculated from the distance between the floor plane and the top plane. For the "
    "length and width I used the detected top part of the box and projected its 3D points into the "
    "coordinate system of the top plane.",
)

add_heading(document, "Short description of the method")
add_body(
    document,
    "Some points in the point cloud are invalid, so I ignored points where z is 0. For RANSAC I used "
    "three random points to make a plane and then counted how many other points are close to that "
    "plane. After the best plane was found, I fitted it again with all inlier points.",
)
add_body(
    document,
    "After the floor was found, I cleaned the mask a little bit with morphological operations. Then I "
    "removed the floor from the point cloud. The remaining points still contain some background, so I "
    "used connected components and the amplitude image to keep the part that most likely belongs to the "
    "box. Then I ran RANSAC again to find the top plane.",
)

add_heading(document, "Weaknesses")
weaknesses = [
    "The result depends quite a lot on the chosen parameters. For example the RANSAC thresholds, the "
    "number of iterations and the mask filtering parameters were selected by trying values on the "
    "given examples.",
    "RANSAC can still find a wrong plane if there is a large flat background area. This happened more "
    "easily near the image border, therefore I removed border-connected components.",
    "The method assumes that the floor and the box top are mostly planar. If the box is partly hidden, "
    "tilted in a strange way or has a bad depth measurement, the plane fit can become worse.",
    "The corner positions depend on the top mask. If this mask misses parts of the box or includes too "
    "many extra pixels, the measured length and width also change.",
    "I used the amplitude image because the box is rather bright in these examples. This is helpful "
    "here, but it would not always work, for example for a dark box or different lighting.",
    "The RANSAC loop is also not the fastest method because many random planes are tested. For these "
    "small examples it is fine, but for larger data it would become slower.",
]
for weakness in weaknesses:
    add_body(document, weakness)

add_heading(document, "Possible improvements")
improvements = [
    "One improvement would be to choose the thresholds automatically from the data instead of setting "
    "them manually.",
    "The second plane could also be forced to be almost parallel to the floor plane, because the top of "
    "the box should normally be parallel to the floor.",
    "The candidate mask could be improved by combining more information, for example distance, "
    "amplitude and maybe local surface normals.",
    "The corner estimation could be made more robust with a better oriented bounding rectangle or by "
    "removing outliers from the top mask before measuring the size.",
    "The final plane fitting could use weights, so points with low amplitude or high residual have less "
    "influence.",
    "For speed, RANSAC could run on fewer sampled points first and then only refine the best result on "
    "the full point cloud.",
    "It would also be useful to print a confidence value, for example based on the number of inliers and "
    "the residual error of the fitted planes.",
]
for improvement in improvements:
    add_body(document, improvement)

document.save(DOCX_PATH)
print(DOCX_PATH)
